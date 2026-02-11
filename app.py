import streamlit as st
import math
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import json
import os
import hashlib

# --- CONFIGURAÇÃO DA PÁGINA ---
# Mudamos o título da aba do navegador aqui
st.set_page_config(page_title="Cálculo de Arc Flash", page_icon="⚡", layout="wide")

# ==============================================================================
# SISTEMA DE AUTENTICAÇÃO E USUÁRIOS
# ==============================================================================
DB_FILE = 'users.json'

def hash_password(password):
    """Gera um hash SHA-256 da senha."""
    return hashlib.sha256(password.encode()).hexdigest()

def load_users():
    """Carrega usuários do arquivo JSON. Se não existir, cria o admin padrão."""
    if not os.path.exists(DB_FILE):
        # Cria admin padrão na primeira vez
        users = {
            "admin": {
                "name": "Administrador",
                "password": hash_password("123"), # Senha padrão inicial
                "role": "admin",
                "approved": True
            }
        }
        save_users(users)
        return users
    try:
        with open(DB_FILE, 'r') as f:
            return json.load(f)
    except:
        return {}

def save_users(users):
    """Salva o dicionário de usuários no JSON."""
    with open(DB_FILE, 'w') as f:
        json.dump(users, f, indent=4)

def login_user(username, password):
    users = load_users()
    if username in users:
        if users[username]['password'] == hash_password(password):
            return users[username]
    return None

def register_user(username, name, password):
    users = load_users()
    if username in users:
        return False, "Usuário já existe!"
    
    users[username] = {
        "name": name,
        "password": hash_password(password),
        "role": "user",
        "approved": False # Padrão: precisa de aprovação
    }
    save_users(users)
    return True, "Cadastro realizado! Aguarde aprovação do administrador."

# ==============================================================================
# APLICAÇÃO PRINCIPAL (CÁLCULO)
# ==============================================================================

# --- FUNÇÃO AUXILIAR DE TEXTO (LATIN-1) ---
def ft(texto):
    try:
        if texto is None: return ""
        return str(texto).encode('latin-1', 'replace').decode('latin-1')
    except:
        return str(texto)

# --- GERADOR DE PDF ---
def gerar_pdf(dados):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, ft('Memorial de Cálculo - Arc Flash'), 0, 1, 'C') # Título PDF Atualizado
    pdf.set_font("Arial", 'I', 9)
    pdf.cell(0, 6, 'Conforme NBR 17227 / IEEE 1584', 0, 1, 'C')
    pdf.ln(4)
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 6, ft(f"Local: {dados['local']}"), 0, 1, 'C')
    eq_texto = dados['eq1']
    if dados['eq2']: eq_texto += f" [{dados['eq2']}]"
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(0, 6, ft(eq_texto), 0, 1, 'C')
    pdf.ln(8)
    pdf.set_fill_color(230, 230, 230)
    pdf.set_font("Arial", 'B', 10)
    pdf.cell(0, 7, ft("1. PARÂMETROS DE ENTRADA"), 1, 1, 'L', 1)
    pdf.set_font("Arial", size=10)
    pdf.ln(2)
    pdf.cell(95, 6, ft(f"Tensão Nominal: {dados['v']:.3f} kV"), 0, 0)
    pdf.cell(95, 6, ft(f"Corrente de Curto (Ibf): {dados['i']:.3f} kA"), 0, 1)
    pdf.cell(95, 6, ft(f"Tempo de Arco: {dados['t']:.4f} s"), 0, 0)
    pdf.cell(95, 6, ft("Configuração: VCB"), 0, 1)
    gap_txt = "(Padrao)" if dados['is_gap_std'] else "(Manual)"
    dist_txt = "(Padrao)" if dados['is_dist_std'] else "(Manual)"
    pdf.cell(95, 6, ft(f"Gap: {dados['g']:.1f} mm {gap_txt}"), 0, 0)
    pdf.cell(95, 6, ft(f"Distância: {dados['d']:.1f} mm {dist_txt}"), 0, 1)
    pdf.ln(5)
    pdf.set_font("Arial", 'B', 10)
    pdf.cell(0, 7, ft("2. ROTEIRO DE CÁLCULO"), 1, 1, 'L', 1)
    pdf.set_font("Courier", size=9)
    pdf.ln(2)
    pdf.cell(0, 5, f"A) Logaritmos:", 0, 1)
    pdf.cell(0, 5, f"   Log(Ibf)={math.log10(dados['i']):.4f} | Log(Gap)={math.log10(dados['g']):.4f}", 0, 1)
    pdf.ln(2)
    pdf.cell(0, 5, ft(f"B) Energia Base (En):"), 0, 1)
    pdf.cell(0, 5, f"   Log(En) = {dados['lg_en']:.4f} -> En = {dados['en_base']:.4f} cal/cm2", 0, 1)
    pdf.ln(2)
    pdf.cell(0, 5, ft(f"C) Fatores:"), 0, 1)
    pdf.cell(0, 5, f"   Tempo ({dados['t']}s/0.2s): {dados['fator_t']:.2f}", 0, 1)
    pdf.cell(0, 5, f"   Distancia (610/{dados['d']})^2: {dados['fator_d']:.3f}", 0, 1)
    pdf.cell(0, 5, f"   Fator Tensao: {dados['fator_v']}", 0, 1)
    pdf.ln(5)
    pdf.set_font("Arial", 'B', 10)
    pdf.cell(0, 7, ft("3. RESULTADO E CLASSIFICAÇÃO"), 1, 1, 'L', 1)
    pdf.ln(3)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, ft(f"Energia Incidente: {dados['e']:.2f} cal/cm²"), 0, 1)
    pdf.set_font("Arial", size=11)
    pdf.set_text_color(0, 0, 0)
    if dados['e'] > 40: pdf.set_text_color(200, 0, 0)
    elif dados['e'] >= 8: pdf.set_text_color(200, 100, 0)
    pdf.cell(0, 8, ft(f"Classificação: {dados['cat']}"), 0, 1)
    pdf.set_text_color(0, 0, 0) 
    pdf.ln(5)
    pdf.set_font("Arial", 'I', 8)
    pdf.cell(0, 5, ft("Nota: A vestimenta deve possuir ATPV superior à energia calculada."), 0, 1)
    return pdf.output(dest='S').encode('latin-1')

# --- GERADOR DE WORD ---
def gerar_word(dados):
    doc = Document()
    head = doc.add_heading('Memorial - Arc Flash', 0) # Título Word Atualizado
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_local = doc.add_paragraph()
    p_local.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l = p_local.add_run(f"Local: {dados['local']}")
    run_l.bold = True
    run_l.font.size = Pt(12)
    eq_texto = dados['eq1']
    if dados['eq2']: eq_texto += f" [{dados['eq2']}]"
    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq = p_eq.add_run(eq_texto)
    run_eq.bold = True
    run_eq.font.size = Pt(11)
    doc.add_paragraph("-" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('1. Parâmetros', level=1)
    p = doc.add_paragraph()
    p.add_run(f"Tensão: {dados['v']:.3f} kV | Corrente: {dados['i']:.3f} kA | Tempo: {dados['t']:.4f} s\n")
    p.add_run(f"Gap: {dados['g']:.1f} mm | Distância: {dados['d']:.1f} mm\n")
    p.add_run("Configuração: VCB")
    doc.add_heading('2. Resultado', level=1)
    p_res = doc.add_paragraph()
    run_res = p_res.add_run(f"{dados['e']:.2f} cal/cm²")
    run_res.bold = True
    run_res.font.size = Pt(16)
    doc.add_paragraph(f"Classificação: {dados['cat']}")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main_app_logic():
    # Título Principal Atualizado
    st.title("⚡ Cálculo de Arc Flash")
    st.markdown(f"Usuário Logado: **{st.session_state['user_name']}**")
    st.markdown("---")

    if 'corrente_stored' not in st.session_state: st.session_state['corrente_stored'] = 17.0
    if 'resultado_icc_detalhe' not in st.session_state: st.session_state['resultado_icc_detalhe'] = None
    if 'ultimo_calculo' not in st.session_state: st.session_state['ultimo_calculo'] = None

    def atualizar_icc():
        try:
            t_kva = st.session_state['k_kva']
            t_v = st.session_state['k_v']
            t_z = st.session_state['k_z']
            usar_motor = st.session_state['k_motor']
            if t_v > 0 and t_z > 0:
                i_nom = (t_kva * 1000) / (math.sqrt(3) * t_v)
                i_cc_trafo = i_nom / (t_z / 100)
                i_motor = 4 * i_nom if usar_motor else 0
                i_total_ka = (i_cc_trafo + i_motor) / 1000
                st.session_state['corrente_stored'] = i_total_ka
                st.session_state['resultado_icc_detalhe'] = {'total': i_total_ka, 'nom': i_nom, 'trafo_ka': i_cc_trafo/1000, 'motor_ka': i_motor/1000}
                st.toast(f"Calculado: {i_total_ka:.3f} kA", icon="✅")
        except: pass

    tab1, tab2 = st.tabs(["🔥 Energia Incidente", "🧮 Icc (Curto)"])

    with tab1:
        st.header("Análise de Energia")
        with st.container(border=True):
            st.subheader("Identificação")
            local_input = st.text_input("Local", placeholder="Ex: Sala Elétrica 01")
            c_eq1, c_eq2 = st.columns(2)
            with c_eq1: eq1_input = st.text_input("Equipamento", placeholder="Ex: QGBT Geral")
            with c_eq2: eq2_input = st.text_input("Detalhe", placeholder="Ex: Disjuntor Entrada")

        st.write("")
        st.info("Parâmetros do Arco:")
        c1, c2, c3 = st.columns(3)
        with c1: tensao = st.number_input("1. Tensão (kV)", value=13.80, format="%.3f")
        with c2: corrente = st.number_input("2. Corrente (kA)", key="corrente_stored", format="%.3f")
        with c3: tempo = st.number_input("3. Tempo (s)", value=0.500, format="%.4f")

        st.caption("Geometria (0 = Padrão)")
        c4, c5 = st.columns(2)
        with c4: gap = st.number_input("Gap (mm)", value=0.0, step=1.0)
        with c5: distancia = st.number_input("Distância (mm)", value=0.0, step=10.0)

        def calcular_completo():
            g_c = gap if gap > 0 else (152.0 if tensao >= 1.0 else 25.0)
            d_c = distancia if distancia > 0 else (914.0 if tensao >= 1.0 else 457.2)
            is_gap_std = (gap <= 0)
            is_dist_std = (distancia <= 0)
            lg_i = math.log10(corrente) if corrente > 0 else 0
            
            if tensao < 1.0:
                k_base, k_i, k_g = -0.555, 1.081, 0.0011
                x_dist = 2.0
                fator_v = 0.85 if tensao < 0.6 else 1.0
            else:
                k_base, k_i, k_g = -0.555, 1.081, 0.0011
                x_dist = 2.0
                fator_v = 1.15

            lg_en = k_base + (k_i * lg_i) + (k_g * g_c)
            en_base = 10 ** lg_en
            fator_t = tempo / 0.2
            fator_d = (610 / d_c) ** x_dist
            e_final = 1.0 * en_base * fator_t * fator_d * fator_v
            
            if e_final < 1.2: cat, cor = "Risco Mínimo", "green"
            elif e_final < 4.0: cat, cor = "Cat 1 / 2", "orange"
            elif e_final < 8.0: cat, cor = "Cat 2", "darkorange"
            elif e_final < 40.0: cat, cor = "Cat 3 / 4", "red"
            else: cat, cor = "PERIGO", "black"

            return {
                'local': local_input, 'eq1': eq1_input, 'eq2': eq2_input,
                'v': tensao, 'i': corrente, 't': tempo, 'g': g_c, 'd': d_c,
                'is_gap_std': is_gap_std, 'is_dist_std': is_dist_std,
                'k_base': k_base, 'k_i': k_i, 'k_g': k_g,
                'lg_en': lg_en, 'en_base': en_base,
                'fator_t': fator_t, 'fator_d': fator_d, 'fator_v': fator_v, 'x_dist': x_dist,
                'e': e_final, 'cat': cat, 'cor': cor
            }

        if st.button("CALCULAR", type="primary", use_container_width=True):
            if tensao > 0 and corrente > 0 and tempo > 0:
                resultado = calcular_completo()
                st.session_state['ultimo_calculo'] = resultado
            else:
                st.warning("Preencha dados obrigatórios.")

        if st.session_state['ultimo_calculo']:
            res = st.session_state['ultimo_calculo']
            st.divider()
            st.markdown(f"**Resultado:** {res['local']} - {res['eq1']}")
            c_res1, c_res2 = st.columns([1, 2])
            c_res1.metric("Energia", f"{res['e']:.2f} cal/cm²")
            c_res2.markdown(f"<div style='background-color:{res['cor']};color:white;padding:15px;text-align:center;border-radius:10px;'><h3>{res['cat']}</h3></div>", unsafe_allow_html=True)
            
            st.write("Downloads:")
            dl1, dl2 = st.columns(2)
            with dl1:
                pdf_data = gerar_pdf(res)
                st.download_button("📥 PDF", data=pdf_data, file_name="memorial.pdf", mime="application/pdf", use_container_width=True)
            with dl2:
                docx_data = gerar_word(res)
                st.download_button("📝 Word", data=docx_data, file_name="memorial.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

    with tab2:
        st.header("Estimativa Curto-Circuito")
        c1, c2 = st.columns(2)
        with c1:
            st.number_input("Potência Trafo (kVA)", value=1000.0, step=100.0, key="k_kva")
            st.number_input("Tensão Sec. (V)", value=380.0, step=10.0, key="k_v")
        with c2:
            st.number_input("Impedância Z (%)", value=5.0, step=0.1, key="k_z")
            st.checkbox("Considerar Motores?", value=True, key="k_motor")
        st.write("")
        st.button("Calcular Icc", on_click=atualizar_icc, type="primary", use_container_width=True)
        dados = st.session_state['resultado_icc_detalhe']
        if dados:
            st.metric("Icc Estimada", f"{dados['total']:.3f} kA")
            st.success("Copiado para Aba 1.")

# ==============================================================================
# LOGIN E CONTROLE DE ACESSO
# ==============================================================================

if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
    st.session_state['user_role'] = None
    st.session_state['user_name'] = None
    st.session_state['user_login'] = None

# Sidebar Admin
def admin_panel():
    st.sidebar.markdown("---")
    st.sidebar.subheader("🛡️ Admin")
    users = load_users()
    
    # Aprovação
    pending = [u for u, data in users.items() if not data['approved']]
    if pending:
        st.sidebar.warning(f"Pendentes: {len(pending)}")
        user_to_approve = st.sidebar.selectbox("Aprovar:", ["..."] + pending)
        if user_to_approve != "...":
            if st.sidebar.button(f"Liberar {user_to_approve}"):
                users[user_to_approve]['approved'] = True
                save_users(users)
                st.sidebar.success("Aprovado!")
                st.rerun()
    else:
        st.sidebar.info("Sem aprovações pendentes.")

    # Exclusão
    st.sidebar.markdown("---")
    all_users = [u for u in users.keys() if u != 'admin' and u != st.session_state['user_login']]
    if all_users:
        user_to_delete = st.sidebar.selectbox("Excluir Usuário:", ["..."] + all_users)
        if user_to_delete != "...":
            if st.sidebar.button(f"Excluir {user_to_delete}"):
                del users[user_to_delete]
                save_users(users)
                st.sidebar.success("Excluído.")
                st.rerun()

# Tela de Login/Registro
if not st.session_state['logged_in']:
    st.title("🔒 Login - Cálculo de Arc Flash")
    
    menu_login = st.selectbox("Opção", ["Entrar", "Criar Conta"])
    
    if menu_login == "Entrar":
        username = st.text_input("Usuário")
        password = st.text_input("Senha", type="password")
        if st.button("Acessar", type="primary"):
            user = login_user(username, password)
            if user:
                if user['approved']:
                    st.session_state['logged_in'] = True
                    st.session_state['user_role'] = user['role']
                    st.session_state['user_name'] = user['name']
                    st.session_state['user_login'] = username
                    st.rerun()
                else:
                    st.warning("⏳ Conta aguardando aprovação do Admin.")
            else:
                st.error("Dados inválidos.")

    elif menu_login == "Criar Conta":
        st.subheader("Novo Cadastro")
        new_user = st.text_input("Defina seu Usuário")
        new_name = st.text_input("Seu Nome")
        new_pass = st.text_input("Defina sua Senha", type="password")
        if st.button("Solicitar Acesso"):
            if new_user and new_name and new_pass:
                success, msg = register_user(new_user, new_name, new_pass)
                if success: st.success(msg)
                else: st.error(msg)
            else:
                st.warning("Preencha tudo.")

else:
    # App Principal Logado
    st.sidebar.success(f"Olá, {st.session_state['user_name']}")
    if st.session_state['user_role'] == 'admin':
        admin_panel()
    
    if st.sidebar.button("Sair"):
        st.session_state['logged_in'] = False
        st.session_state['user_role'] = None
        st.rerun()

    main_app_logic()
